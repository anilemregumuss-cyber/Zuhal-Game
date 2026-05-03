from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Sayfa kenar boşlukları
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(text, level=1, color="1a1a1a", space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    if level == 1:
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C8A400')
        pBdr.append(bottom)
        pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    size = {1: 16, 2: 13, 3: 11}
    run.font.size = Pt(size.get(level, 11))
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_para(text, bold=False, color="333333", size=10, space_before=2, space_after=2, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_checkbox(text, bold_part=None, note=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.3)
    run_box = p.add_run("☐  ")
    run_box.font.size = Pt(10)
    run_box.font.color.rgb = RGBColor.from_string("C8A400")
    if note:
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string("666666")
        run.italic = True
    elif bold_part:
        idx = text.find(bold_part)
        if idx >= 0:
            run1 = p.add_run(text[:idx])
            run1.font.size = Pt(10)
            run1.font.color.rgb = RGBColor.from_string("222222")
            run_b = p.add_run(bold_part)
            run_b.bold = True
            run_b.font.size = Pt(10)
            run_b.font.color.rgb = RGBColor.from_string("222222")
            rest = text[idx+len(bold_part):]
            if rest:
                run2 = p.add_run(rest)
                run2.font.size = Pt(10)
                run2.font.color.rgb = RGBColor.from_string("222222")
        else:
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("222222")
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("222222")
    return p

def add_warning(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("⚠  " + text)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("B84000")
    return p

def add_section_box(title, color_hex="F5F0DC"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("  " + title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1a1a1a')
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0)
    return p

# ── BAŞLIK ──────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(4)
run = p_title.add_run("ZUHAL MÜZİK")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor.from_string("C8A400")

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(2)
run2 = p_sub.add_run("HALFTİME 50 KAMPANYA KILAVUZU")
run2.bold = True
run2.font.size = Pt(15)
run2.font.color.rgb = RGBColor.from_string("1a1a1a")

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_desc.paragraph_format.space_before = Pt(0)
p_desc.paragraph_format.space_after = Pt(12)
run3 = p_desc.add_run("Mağaza Operasyon El Kitabı")
run3.italic = True
run3.font.size = Pt(10)
run3.font.color.rgb = RGBColor.from_string("888888")

# ── MAĞAZA TAKVİMİ ──────────────────────────────────────
add_heading("MAĞAZA TAKVİMİ", level=1, color="1a1a1a")

stores = [
    ("1","15 Mayıs","Cuma","Zuhal Akasya"),
    ("2","22 Mayıs","Cuma","Ankara Kentpark"),
    ("3","31 Mayıs","Pazar","Zuhal Hilltown"),
    ("4","5 Haziran","Cuma","Bursa Downtown"),
    ("5","12 Haziran","Cuma","Zuhal UNIQ"),
    ("6","19 Haziran","Cuma","İzmir Basmane"),
    ("7","26 Haziran","Cuma","Zuhal Kanyon"),
    ("8","4 Temmuz","Cuma","İzmir Mavibahçe"),
    ("9","11 Temmuz","Cumartesi","Antalya Zuhal"),
    ("10","18 Temmuz","Cumartesi","Adana 01 Burda"),
    ("11","25 Temmuz","Cumartesi","Temaworld Zuhal"),
    ("12","1 Ağustos","Cumartesi","Bodrum Midtown"),
    ("13","8 Ağustos","Cumartesi","Ankara Kızılay"),
    ("14","15 Ağustos","Cumartesi","Zuhal Tünel  ★  FİNALE"),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
hdr = table.rows[0].cells
headers = ["#", "TARİH", "GÜN", "MAĞAZA"]
widths = [Cm(1.2), Cm(3.5), Cm(3.5), Cm(8)]
for i, (cell, htext) in enumerate(zip(hdr, headers)):
    cell.width = widths[i]
    set_cell_bg(cell, "1a1a1a")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(htext)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("C8A400")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for row_data in stores:
    row = table.add_row()
    is_finale = row_data[0] == "14"
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        cell.width = widths[i]
        if is_finale:
            set_cell_bg(cell, "1a1a1a")
        elif int(row_data[0]) % 2 == 0:
            set_cell_bg(cell, "F8F8F8")
        else:
            set_cell_bg(cell, "FFFFFF")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(9.5)
        run.bold = is_finale
        run.font.color.rgb = RGBColor.from_string("C8A400" if is_finale else "222222")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── 3 HAFTA ÖNCE ────────────────────────────────────────
add_section_box("3 HAFTA ÖNCE — Kondisyon Raporu")
add_checkbox("Mağazadaki kondisyonu düşük ürünlerin listesi hazırlanır", bold_part="kondisyonu düşük ürünlerin")
add_para("       (çizik, hatalı, eksik parçalı ürünler)", size=9, color="666666", space_before=0)
add_checkbox("Liste Anıl Emre Gümüş'e iletilir", bold_part="Anıl Emre Gümüş'e")

# ── 2 GÜN ÖNCE ──────────────────────────────────────────
add_section_box("KAMPANYA GÜNÜNDEN 2 GÜN ÖNCE")
add_heading("Merkezden Mağazaya Ulaşacaklar", level=2, color="444444", space_before=8)
add_checkbox("Kiosk ekranı")
add_checkbox("Kampanya totemi (kondisyonlu ürünler için)")
add_checkbox("Tavana asılacak Halftime görselleri")
add_checkbox("Halftime kampanya ürün etiketleri")
add_checkbox("Kampanya ürün listesi ve indirim oranları")

add_heading("Mağazanın Yapacakları", level=2, color="444444", space_before=8)
add_checkbox("Personel kampanya ürünlerini inceler, bilgilerini tazeler", bold_part="bilgilerini tazeler")
add_checkbox("Kampanya ürünleri ön plana çıkarılmaya başlanır")

add_heading("Dijital Ekranlar (varsa)", level=3, color="666666", space_before=6)
add_checkbox("Pazartesi: kampanya görselleri yüklenir — \"Bu Cuma/Cumartesi [Mağaza]'ndayız\" duyurusu")
add_checkbox("Kampanya gününden sonraki 2 gün: kampanya videoları yayınlanır")
add_checkbox("Sonraki Pazartesi: yeni mağaza duyurusu ile devam edilir")

# ── 1 GÜN ÖNCE ──────────────────────────────────────────
add_section_box("KAMPANYA GÜNÜNDEN 1 GÜN ÖNCE — Mağaza Kapanışında")
add_checkbox("Tüm kampanya ürünleri Halftime etiketleriyle etiketlenir", bold_part="Halftime etiketleriyle")
add_checkbox("Tavan görselleri yerleştirilir, dekorasyon tamamlanır")
add_checkbox("Totem kurulur, kondisyonlu ürünler toteme yerleştirilir")
add_checkbox("Kiosk kasanın yanına kurulur ve test edilir", bold_part="kasanın yanına")
add_checkbox("Yoğunluğa göre diğer mağazalardan ek personel planlanır")

# ── KAMPANYA GÜNÜ ────────────────────────────────────────
add_section_box("KAMPANYA GÜNÜ — Saat 10:00'da Hazır")
add_heading("Açılış Kontrol Listesi", level=2, color="444444", space_before=8)
add_checkbox("Tüm ürünler etiketli ve ön planda")
add_checkbox("Mağaza Halftime görselleriyle süslenmiş")
add_checkbox("Kiosk çalışır durumda, başında görevli personel var")
add_checkbox("Tüm personel kampanya ürünlerini ve indirim oranlarını biliyor")

add_heading("Gün Boyunca", level=2, color="444444", space_before=8)
add_checkbox("Kampanya ürünü satın alan her müşteri oyuna yönlendirilir")
add_checkbox("Her müşterinin video çekip @zuhalmuzik etiketleyerek paylaşması sağlanır", bold_part="@zuhalmuzik etiketleyerek")
add_checkbox("Kiosk görevlisi oyun akışını yönetir")
add_warning("Sosyal medya paylaşımı kritiktir — kampanyadan sonraki 2 gün bu içerikler Zuhal hesabından paylaşılacak. Müşteri paylaşımı olmadan kampanya sosyal etki yaratamaz.")

# ── KAMPANYA SONRASI ─────────────────────────────────────
add_section_box("KAMPANYA SONRASI — Akşam Kapanışında")
add_checkbox("Tüm Halftime ürün etiketleri sökülerek toplanır")
add_checkbox("Tavan görselleri toplanır")
add_checkbox("Totem ve kiosk sökülerek paketlenir")
add_checkbox("Tüm malzemeler bir sonraki mağazaya gönderilmek üzere hazırlanır", bold_part="bir sonraki mağazaya")
add_warning("Kampanya materyalleri (totem, tavan görselleri, etiketler, kiosk) tek kullanımlık değildir. 14 mağazanın tamamında kullanılacak şekilde kaliteli üretilmiştir. Hasarlı kullanım kabul edilmez.")

# ── KİOSK KULLANIM ───────────────────────────────────────
add_section_box("KİOSK KULLANIM KURALI")
add_para(
    "Kiosk kasanın yanında konumlanır. Kampanya kapsamında Halftime ürünü satın almış müşteriler oyuna katılabilir. "
    "Doğrulama kiosk görevlisi tarafından yapılır — müşteriden ek bilgi istenmez. "
    "Oyuncu adını ve soyadını kendisi girer; kiosk görevlisi oyunu başlatır.",
    size=10, color="333333", space_before=6, space_after=4
)

# ── OYUN KURGUSU ─────────────────────────────────────────
add_section_box("OYUN KURGUSU — Kiosk Görevlisi Bilgi Notu")
add_para("Her oyuncunun 2 oyun hakkı vardır:", bold=True, size=10, color="222222", space_before=6)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
r1 = p.add_run("1. Oyun: ")
r1.bold = True
r1.font.size = Pt(10)
r1.font.color.rgb = RGBColor.from_string("C8A400")
r2 = p.add_run("Deneme turu — ödül verilmez, sonuç gösterilir")
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor.from_string("444444")

p2 = doc.add_paragraph()
p2.paragraph_format.left_indent = Cm(0.5)
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after = Pt(4)
r3 = p2.add_run("2. Oyun: ")
r3.bold = True
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor.from_string("C8A400")
r4 = p2.add_run("Gerçek tur — çıkan sonuç oyuncunun ödülüdür")
r4.font.size = Pt(10)
r4.font.color.rgb = RGBColor.from_string("444444")

add_heading("Ödül Tablosu", level=2, color="444444", space_before=8)

prize_table = doc.add_table(rows=1, cols=3)
prize_table.style = 'Table Grid'
prize_hdr = prize_table.rows[0].cells
for cell, txt in zip(prize_hdr, ["SONUÇ", "ZAMANLAMA", "ÖDÜL"]):
    set_cell_bg(cell, "333333")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(txt)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("C8A400")

prizes = [
    ("🏆  Mükemmel", "Tom'dan 50ms erken", "%15 İndirim"),
    ("🎯  Harika", "Tom'dan 130ms erken\nveya 0–20ms geç", "%10 İndirim"),
    ("🎵  İyi", "Tom'dan 280ms erken", "2 Saat Stüdyo Prova"),
    ("🎓  İdare Eder", "Tom'dan 500ms erken", "Akademi 1 Ders / 1 Saat Stüdyo / %5 İndirim (rastgele)"),
    ("🎁  Bez Çanta", "20ms+ geç veya hiç basmama", "50. Yıl Bez Çanta"),
]
for i, (sonuc, zaman, odul) in enumerate(prizes):
    row = prize_table.add_row()
    bg = "F8F8F8" if i % 2 == 0 else "FFFFFF"
    for cell, txt in zip(row.cells, [sonuc, zaman, odul]):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(txt)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("222222")

add_para("Oyuncu ödülünü mağaza ekibine ad soyad ve kupon kodu ile birlikte gösterir.",
         size=9, color="666666", space_before=6)

doc.save("/home/user/Zuhal-Game/Halftime50_Kampanya_Kilavuzu.docx")
print("Dosya oluşturuldu.")
